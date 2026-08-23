#!/usr/bin/env python3
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SKILL_SCRIPTS = Path(
    "/Users/abirzu/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.802.11031/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


SOURCE = Path(".tmp/axian-bom/AXIAN-OCI-Observability-Management-RFP-Answers-Latest.docx")
OUTPUT = Path(".tmp/axian-bom/AXIAN-OCI-Observability-Management-RFP-Answers-Latest-BOM.docx")


PLATFORM_ROWS = [
    (
        "Oracle Communications Unified Assurance Hyperscale core",
        "Existing / Assurance dependency",
        "One Group logical platform with production, non-production and DR capacity sized by event rate, topology size, service models, concurrent operators and OPCO isolation.",
        "Supplied and licensed by the Assurance workstream; not priced in the OCI O&M estimate.",
    ),
    (
        "Unified Assurance collectors and source adapters",
        "Required dependency",
        "Collector capacity per OPCO, failure domain and source type; count from the approved telecom, infrastructure, OEM and application interface catalogue.",
        "Assurance/product teams confirm supported collectors; uncovered proprietary interfaces require CSS/partner work.",
    ),
    (
        "Unified Assurance native ServiceNow Adapter",
        "Required interface",
        "One Group ServiceNow integration plus non-production test path; validate 50-300 tickets/day per OPCO, retries, idempotency and bidirectional state.",
        "Verify UA adapter entitlement. This remains the only incident create/update/close path; no OIC SKU is included.",
    ),
    (
        "Oracle Enterprise Manager Cloud Control, repository and Management Agents",
        "Existing / conditional",
        "Reuse the approved OEM estate; validate OMS/repository HA, target count, plug-ins, agent versions, API/export load and DR.",
        "Operations involvement (Pablo) is required. New OEM infrastructure is not assumed until discovery identifies a gap.",
    ),
    (
        "OEM Diagnostics, Tuning and other applicable Management Packs",
        "Conditional",
        "Quantity follows the applicable license metric and managed-target scope. Enable only the pack features selected by the database/OEM team.",
        "Entitlement must be verified before use; no pack is assumed merely because OEM is installed.",
    ),
    (
        "OCI Monitoring",
        "Required OCI O&M",
        "Compartments, custom metric streams, dimensions, alarms, notification fan-out and query volume; separate Group and OPCO scopes.",
        "Use the current Oracle rate card/Cost Estimator after metric cardinality and alarm design are known.",
    ),
    (
        "OCI Logging",
        "Required OCI O&M",
        "Daily service/custom log ingestion, active search retention, log groups and cross-region/residency routing.",
        "Product teams own log quality; exclude unapproved raw payloads and duplicate ingestion.",
    ),
    (
        "Connector Hub",
        "Required OCI O&M",
        "Service connectors per source/destination/residency route, including Log Analytics, Object Storage, Streaming, Notifications and Functions.",
        "Validate delivery semantics, duplicate handling, dead-letter/replay design and route-health monitoring.",
    ),
    (
        "OCI Log Analytics",
        "Required OCI O&M",
        "GB/day of selected analytical logs, searchable retention, entities, sources, parsers, scheduled tasks, dashboards and concurrent users.",
        "Do not size from the full 2 TB/day figure without filtering and tiering. LoganAI/model usage is assessed separately.",
    ),
    (
        "Log Analytics Kubernetes Monitoring Solution",
        "Required for approved OKE clusters",
        "One deployment per cluster; size by clusters, nodes, pods, namespaces, log volume, object-state volume and retention.",
        "Private OKE clusters use the supported manual deployment path; cluster access remains with the platform team.",
    ),
    (
        "LoganAI",
        "Conditional",
        "Enabled users, investigation frequency, selected logs/metrics and model calls after realm/region and residency validation.",
        "Analyst assistance only; not the dedicated incident commander. Human validation and GenAI consumption apply.",
    ),
    (
        "OCI Application Performance Monitoring",
        "Required for priority applications",
        "APM domains by isolation model; size by application servers, spans, browser sessions, synthetic monitors/vantage points and retention.",
        "Application teams instrument Java, .NET, PHP and APIs and provide service IDs, journeys and SLOs.",
    ),
    (
        "OCI Management Agent and Management Gateway",
        "Required where hybrid collection applies",
        "Agents per supported target and redundant gateways per disconnected/private network zone; include non-production and DR.",
        "Distinct from OEM Management Agents. Validate proxy, ports, certificates, upgrade and failure-recovery procedures.",
    ),
    (
        "OCI Database Management",
        "Conditional",
        "Managed database/host count, deployment type, collection method, retention and administrative scope.",
        "Database team supplies connectivity, credentials, privileges and commercial approval.",
    ),
    (
        "OCI Operations Insights",
        "Conditional",
        "Enabled database/host resources, historical capacity data, SQL insights scope and forecast horizon.",
        "Database/OEM team validates prerequisites and operational use of findings.",
    ),
    (
        "OCI Notifications and Events",
        "Required OCI O&M",
        "Topics, subscriptions, event rules, endpoints, retry policy and regional/OPCO fan-out.",
        "Used for governed notification and automation entry points; UA remains the incident-policy authority.",
    ),
    (
        "OCI Object Storage and Archive Storage",
        "Required OCI O&M",
        "Raw/evidence GB per day, 6-12 month retention by data class, lifecycle tiers, replay frequency, replication and retrieval expectations.",
        "Apply residency, encryption, legal-hold and lifecycle rules before costing.",
    ),
    (
        "OCI Streaming and/or Queue",
        "Conditional",
        "Partition/throughput, message size, retention, consumer groups, backlog and replay requirements for routes needing buffering.",
        "Include only where measured throughput or decoupling requirements exceed direct connector/function patterns.",
    ),
    (
        "OCI Generative AI project, Responses API and model inference",
        "Required for the custom incident agent",
        "One project per approved isolation/environment pattern; size by incidents, turns, input/output tokens, reasoning level, concurrency and one-month-plus evidence context.",
        "OCI AI Team involvement (Alex Negrea) is required. Model/region, quota, retention, evaluation and residency must be accepted.",
    ),
    (
        "OCI Generative AI Conversations, Files and Vector Stores",
        "Conditional agent capability",
        "Conversation count/retention, runbook and evidence corpus size, vector-store growth, file ingestion and retrieval calls.",
        "Include only the approved memory/RAG pattern; source provenance, deletion and access isolation are acceptance gates.",
    ),
    (
        "OCI API Gateway",
        "Required for the governed tool plane",
        "Private/approved public endpoints, requests, payload size, rate limits, custom domains, certificates and environment/region count.",
        "Exposes only allowlisted UA, OCI, OEM/vendor, ServiceNow/CMDB-read and customer API contracts.",
    ),
    (
        "OCI Functions",
        "Required for bounded API adapters",
        "Function count, invocations, execution time, memory, concurrency, network access, retries and deployment environments.",
        "CSS/partner develops schema validation, redaction and adapters. Read-only is default; mutations require approval and rollback evidence.",
    ),
    (
        "OCI IAM, Vault, Audit and Cloud Guard",
        "Required security baseline",
        "Compartments, groups/dynamic groups, policies, secrets/keys, rotation, audit retention, detector/target scope and OPCO isolation.",
        "Observability & Security defines controls; product teams request least-privilege access to their resources.",
    ),
    (
        "Private networking, DNS, certificates and WAF where public exposure is approved",
        "External dependency / conditional",
        "Connectivity per OPCO and region, private endpoints, DRG/VPN/FastConnect capacity, DNS zones, certificates and WAF policies.",
        "Network/security product teams provide these components. They are dependencies, not part of the OCI O&M service estimate.",
    ),
]


WORK_ROWS = [
    (
        "Discovery and detailed design",
        "Source inventory, volume/cardinality study, retention/residency classes, target architecture, security model and final commercial BOM.",
        "All OPCOs and priority tools sampled; 2 TB/day and ticket-volume assumptions reconciled.",
    ),
    (
        "OCI observability foundation",
        "Compartments, IAM, Vault, log groups, Monitoring, Logging, Connector Hub, Notifications, Object Storage and monitor-the-monitoring controls.",
        "Infrastructure-as-code, least privilege, operational alarms and non-production validation.",
    ),
    (
        "Source onboarding and normalization",
        "Supported agents/collectors, Log Analytics parsers, canonical IDs, enrichment, quality checks, replay and UA forwarding.",
        "Priority interface catalogue; completeness, latency, duplicate and loss tests.",
    ),
    (
        "OKE and application observability",
        "Kubernetes Monitoring Solution, APM/OpenTelemetry instrumentation, synthetic journeys, dashboards and SLO alarms.",
        "Approved clusters and priority applications; trace continuity and failure-injection evidence.",
    ),
    (
        "OEM and native ServiceNow integration",
        "OEM export/context mapping into UA and UA native ServiceNow Adapter workflow validation.",
        "Operations (Pablo), OEM and ITSM acceptance; entitlement, idempotency, retry and outage-recovery tests.",
    ),
    (
        "Dashboards, alerts and runbooks",
        "Group/OPCO views, alarm standards, suppression, maintenance windows, escalation policy and evidence links.",
        "Versioned KPI/SLO catalogue and signed use-case acceptance.",
    ),
    (
        "Dedicated incident-commander agent",
        "Responses API coordinator, governed RAG, API Gateway/Functions tools, citations, evidence package, guardrails, evaluation and approval UX.",
        "OCI AI Team (Alex Negrea), CSS/qualified partner and Security acceptance; read-only PoV before any mutation discussion.",
    ),
    (
        "PoV, resilience and handover",
        "Load, scale, HA/DR, security, data-residency and restore tests; runbooks, training, support and rollback package.",
        "Signed acceptance metrics, four-hour RTO, five-to-fifteen-minute RPO and support/escalation readiness.",
    ),
]


def insert_paragraph_before(doc, target, text="", style=None):
    paragraph = doc.add_paragraph(text, style=style)
    target._p.addprevious(paragraph._p)
    return paragraph


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def format_table(table, widths, status_column=None):
    table.style = "List Table 3"
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=110,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 110, "end": 110},
    )
    repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if status_column is not None and col_index == status_column and row_index > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)
            if row_index > 0 and col_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def add_table_before(doc, target, headers, rows, widths, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    format_table(table, widths, status_column=status_column)
    target._p.addprevious(table._tbl)
    return table


def main():
    doc = Document(SOURCE)

    toc_reference = next(
        p for p in doc.paragraphs
        if p.text.strip() == "7. Official Oracle references" and p.style.name == "Normal"
    )
    toc_reference.text = "7. Solution bill of materials (BOM)"
    toc_new_reference = doc.add_paragraph("8. Official Oracle references", style="Normal")
    toc_reference._p.addnext(toc_new_reference._p)

    references_heading = next(
        p for p in doc.paragraphs
        if p.text.strip() == "7. Official Oracle references" and p.style.name.startswith("Heading")
    )
    references_heading.text = "8. Official Oracle references"

    heading = insert_paragraph_before(doc, references_heading, "7. Solution bill of materials (BOM)", "Heading 1")
    heading.paragraph_format.page_break_before = True
    lead = insert_paragraph_before(
        doc,
        references_heading,
        "This BOM is a solution and procurement baseline, not a binding commercial quote. Exact service quantities and SKUs must be generated from measured discovery data and the applicable Oracle agreement. The current 2 TB/day figure is a planning input only until it is split by OPCO, region, telemetry type, searchable retention and archive retention.",
    )
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    insert_paragraph_before(doc, references_heading, "Platform and subscription BOM", "Heading 2")
    add_table_before(
        doc,
        references_heading,
        ["Layer / component", "Status", "Baseline quantity or sizing driver", "Commercial / implementation note"],
        PLATFORM_ROWS,
        [2260, 1600, 3020, 2588],
        status_column=1,
    )

    insert_paragraph_before(doc, references_heading, "Reference configuration assumptions", "Heading 2")
    assumptions = [
        "Production and non-production are required; DR resources are duplicated in a second approved region only after service availability, residency and the stated RTO/RPO are validated.",
        "Retention is 6-12 months by data class, with searchable, analytical and archive tiers; at least one month of approved history is planned for AI use cases.",
        "All commercial quantities remain TBD after discovery; the Oracle Cloud Cost Estimator is evaluation guidance and the applicable Oracle quote/rate card is authoritative.",
    ]
    for text in assumptions:
        insert_paragraph_before(doc, references_heading, text, "Small bullet")

    insert_paragraph_before(doc, references_heading, "Implementation and professional-services BOM", "Heading 2")
    add_table_before(
        doc,
        references_heading,
        ["Work package", "Deliverable", "Sizing / acceptance basis"],
        WORK_ROWS,
        [2260, 4100, 3108],
    )

    last_reference = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("OCI Functions overview:")
    )
    for text in [
        "Enterprise AI Agents in OCI Generative AI: https://docs.oracle.com/en-us/iaas/Content/generative-ai/agents.htm",
        "Estimating OCI monthly costs: https://docs.oracle.com/en-us/iaas/Content/Billing/Tasks/signingup_topic-Estimating_Costs.htm",
    ]:
        new_reference = doc.add_paragraph(text, style="Small bullet")
        last_reference._p.addnext(new_reference._p)
        last_reference = new_reference

    # Keep the expanded source list on one page so the Oracle legal closing
    # page is not split for only a few trailing lines.
    for paragraph in doc.paragraphs:
        if "https://docs.oracle.com" in paragraph.text:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(8.4)

    for paragraph in doc.paragraphs:
        if "July, 2026, Version 1.0" in paragraph.text:
            paragraph.runs[0].text = "August, 2026, Version 1.3"
            for run in paragraph.runs[1:]:
                run.text = ""

    processed_footer_parts = set()
    for section in doc.sections:
        footer = section.footer
        part_key = str(footer.part.partname)
        if part_key in processed_footer_parts:
            continue
        processed_footer_parts.add(part_key)
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                if "Version 1.2" in run.text:
                    run.text = run.text.replace("Version 1.2", "Version 1.3")

    doc.core_properties.version = "1.3"
    doc.core_properties.comments = (
        "Added solution, subscription and professional-services BOM with discovery-based sizing and licensing boundaries."
    )
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
