from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


SRC = Path(".tmp/axian-oci-only/current.docx")
OUT = Path(".tmp/axian-oci-only/AXIAN-OCI-Observability-Management-RFP-Answers-Latest.docx")

LINKS = {
    "monitoring": "https://docs.oracle.com/en-us/iaas/Content/Monitoring/home.htm",
    "logging": "https://docs.oracle.com/en-us/iaas/Content/Logging/home.htm",
    "connector": "https://docs.oracle.com/en-us/iaas/Content/connector-hub/overview.htm",
    "logan": "https://docs.oracle.com/en-us/iaas/log-analytics/home.htm",
    "k8s": "https://docs.oracle.com/en-us/iaas/log-analytics/doc/kubernetes-solution.html",
    "loganai": "https://docs.oracle.com/en-us/iaas/log-analytics/doc/use-loganai.html",
    "apm": "https://docs.oracle.com/en-us/iaas/application-performance-monitoring/home.htm",
    "agent": "https://docs.oracle.com/en-us/iaas/management-agents/home.htm",
    "notifications": "https://docs.oracle.com/en-us/iaas/Content/Notification/home.htm",
    "pricing_api_doc": "https://docs.oracle.com/en-us/iaas/Content/Billing/Tasks/signingup_topic-Estimating_Costs.htm#accessing_list_pricing",
    "pricing_api": "https://apexapps.oracle.com/pls/apex/cetools/api/v1/products/",
    "cost_estimator": "https://www.oracle.com/cloud/costestimator.html",
}


def remove_row(table, row):
    table._tbl.remove(row._tr)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_hyperlink(paragraph, text, url, *, bold=False):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_text(cell, text, *, bold=False):
    p = clear_cell(cell)
    r = p.add_run(text)
    r.bold = bold
    return p


def set_linked_cell(cell, label, url, description=""):
    p = clear_cell(cell)
    add_hyperlink(p, label, url, bold=True)
    if description:
        p.add_run(" — " + description)


def add_row(table, values):
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_text(cell, value)
    return row


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    for i, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, "312D2A")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(7.5)


def find_table(document, first_header):
    for table in document.tables:
        if table.rows and table.rows[0].cells and " ".join(table.rows[0].cells[0].text.split()) == first_header:
            return table
    raise KeyError(f"Table not found: {first_header}")


def replace_para_text(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_architecture_table_after(paragraph):
    table = paragraph._parent.add_table(rows=1, cols=4, width=Inches(8.1))
    for i, title in enumerate(["Telemetry source", "Manual / open collection", "OCI Observability data plane", "Operator outcome and hand-off"]):
        set_text(table.rows[0].cells[i], title, bold=True)
    flows = [
        (
            "Customer-managed Kubernetes\nNodes, pods, namespaces, audit and workload logs; Kubernetes object state",
            "Manual Log Analytics Kubernetes Monitoring Solution deployment using the documented Helm/manifests path. Validate cluster version, RBAC, proxy/private egress, certificates, Management Agent/Gateway reachability, upgrades and rollback.",
            "Log Analytics stores and analyzes the selected Kubernetes logs and object state. Monitoring holds approved custom/platform metrics. No OKE control-plane dependency is assumed.",
            "Cluster/workload/node/pod views, searches, dashboards and alerts. Platform teams remain accountable for Kubernetes access, capacity, upgrades and remediation.",
        ),
        (
            "Applications and APIs\nPriority services, customer journeys and middleware",
            "APM agents and OpenTelemetry instrumentation. Application teams supply service identity, trace propagation, approved attributes, synthetic journeys and SLOs.",
            "OCI APM supplies traces, service maps, browser/API visibility and synthetics. The commercial working assumption meters APM active storage through B95634 at one 300 GB Log Analytics Active Storage Unit.",
            "Trace continuity, dependency evidence, journey health and SLO alarms; evidence is exposed by documented APIs to approved external consumers.",
        ),
        (
            "OCI and supported hybrid logs/metrics\nService logs, custom logs and custom metrics",
            "OCI Logging, supported Management Agent/Gateway collection and Connector Hub routes. Collection is limited to approved data classes and supported sources.",
            "Logging retains searchable logs; Log Analytics parses, enriches and investigates selected analytical logs; Monitoring evaluates metrics and alarms; Notifications distributes approved alarms.",
            "Dashboards, alerts, LoganAI-assisted investigation and evidence links. External ticketing, assurance and product-remediation workflows are outside this document and BOM.",
        ),
    ]
    for vals in flows:
        add_row(table, vals)
    format_table(table, [1.55, 2.25, 2.25, 2.15])
    paragraph._p.addnext(table._tbl)
    return table


doc = Document(SRC)

# Version and scope metadata.
for section in doc.sections:
    for container in (section.header, section.footer):
        for p in container.paragraphs:
            for run in p.runs:
                run.text = run.text.replace("Version 1.3", "Version 1.4")

for p in doc.paragraphs:
    for run in p.runs:
        run.text = run.text.replace("Version 1.3", "Version 1.4")

# Front matter and table of contents.
replace_para_text(doc.paragraphs[8],
    "This scoped response contains only the RFP answers and bill of materials delivered by the Observability & Security team through OCI Observability services. The target environment uses customer-managed Kubernetes, so no managed Kubernetes service is assumed. Cross-product assurance, ITSM, database-management, networking, security-platform and AI-agent implementation are outside this document and are referenced only as external interfaces where necessary.")
front_replacements = {
    18: "One external ITSM integration exists, but its product, licensing, workflow and operations are outside this OCI Observability response.",
    19: "OCI Log Analytics Kubernetes Monitoring Solution is proposed for customer-managed Kubernetes using the documented manual Helm/manifests deployment path and platform-specific validation.",
    22: "No Operations/Pablo involvement is required by this scoped OCI Observability response. Product teams own their platforms, access, remediation and external workflows.",
    26: "4. Customer-managed Kubernetes monitoring answer",
    29: "7. OCI Observability bill of materials (BOM)",
    30: "8. Implementation documentation and supporting references",
}
for idx, text in front_replacements.items():
    replace_para_text(doc.paragraphs[idx], text)

# Executive summary and scope boundary.
replace_para_text(doc.paragraphs[33],
    "Scope of this extract: only responses materially delivered by OCI Observability services and the Observability & Security team. Assurance, ITSM, database-management, network/security platform and other product components are excluded from the solution BOM and implementation scope. Product teams remain accountable for their platforms, access, instrumentation, remediation and external workflows. OCI AI Team (Alex Negrea) owns any optional dedicated GenAI incident coordinator. Status convention: Compliant means a documented OCI service supplies the stated capability when configured; Partial means discovery, custom telemetry, product-team input, sizing or PoV evidence is still required.")
exec_summary = {
    35: "OCI Observability provides the central telemetry, analysis and alerting layer for the application, API, Kubernetes and supported hybrid visibility gaps identified in the RFP.",
    36: "OCI Monitoring supplies metrics, alarms and independent service-health indicators for OCI resources and approved custom telemetry.",
    37: "OCI Logging and Connector Hub ingest, retain and route approved service and custom logs without introducing a separate ticketing path.",
    38: "OCI Log Analytics supplies parsing, enrichment, clustering, investigation, dashboards and LoganAI-assisted analysis for selected analytical logs.",
    39: "OCI Log Analytics Kubernetes Monitoring Solution is proposed for customer-managed Kubernetes. Deployment is manual through the documented Helm/manifests pattern and requires cluster-specific validation; no OKE service is assumed.",
    40: "OCI APM and OpenTelemetry supply application/API traces, service topology, browser visibility and synthetic journeys. For this proposal, APM active storage is commercially presumed to use OCI Log Analytics Active Storage SKU B95634 at one Storage Unit per 300 GB active storage; the final Oracle quote must confirm the mapping.",
    41: "OCI Management Agent and Management Gateway provide supported hybrid collection paths where required and are distinct from any external product agent.",
    42: "OCI Notifications distributes approved alarms. External assurance, ITSM/ticketing, database-management and product-remediation workflows are outside this response.",
    43: "LoganAI assists analysts inside Log Analytics but is not an agentic incident commander. An optional dedicated GenAI coordinator is an external OCI AI Team workstream and is not included in this Observability BOM.",
    44: "All production sizing remains subject to measured ingestion, retention, metric cardinality, cluster inventory, trace volume, regions, residency and non-production/DR requirements.",
    45: "The PoV must demonstrate ingestion completeness and latency, manual Kubernetes deployment, parser quality, trace continuity, alert delivery, role isolation, resilience and rollback before production acceptance.",
}
for idx, text in exec_summary.items():
    replace_para_text(doc.paragraphs[idx], text)

# Remove the old cross-product architecture images and narrative, then insert a scoped native Word architecture.
arch_heading = doc.paragraphs[48]
for idx in [55, 54, 53, 52, 51, 50, 49]:
    delete_paragraph(doc.paragraphs[idx])
arch_table = insert_architecture_table_after(arch_heading)
p = OxmlElement("w:p")
arch_table._tbl.addnext(p)
genai_p = doc.add_paragraph()
genai_p._p.getparent().remove(genai_p._p)
p.addnext(genai_p._p)
genai_p.style = doc.styles["Normal"]
replace_para_text(genai_p,
    "Optional GenAI correlation extension — outside the OCI Observability BOM: an OCI AI Team-owned coordinator may consume curated Monitoring, Logging, Log Analytics and APM evidence through documented, allowlisted APIs. It must preserve source provenance, enforce least privilege and remain read-only until separately designed, evaluated and approved. It does not replace OCI telemetry services or own any external incident-write workflow.")
caption_p = doc.add_paragraph()
caption_p._p.getparent().remove(caption_p._p)
genai_p._p.addnext(caption_p._p)
replace_para_text(caption_p,
    "Figure 1. Scoped OCI Observability architecture and communication flows. Arrows are represented left-to-right by table columns; external products and their operations are intentionally excluded.")
caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in caption_p.runs:
    r.italic = True
    r.font.size = Pt(8)

# Remaining section headings and Kubernetes answer.
for p in doc.paragraphs:
    t = " ".join(p.text.split())
    if t == "4. OKE monitoring answer":
        replace_para_text(p, "4. Customer-managed Kubernetes monitoring answer")
    elif t.startswith("The RFP explicitly includes cloud/container platforms"):
        replace_para_text(p,
            "The RFP explicitly includes containers, application platforms, central logs/metrics/traces and topology-aware operations. OCI Log Analytics Kubernetes Monitoring Solution is therefore proposed in the OCI Observability scope for the customer's Kubernetes estate. For each approved cluster, the solution collects selected Kubernetes logs, metrics and object state and exposes cluster, workload, node and pod views. It is combined with OCI Monitoring for approved metrics and alarms, OCI Logging for service/custom logs, Connector Hub for routing and OCI APM/OpenTelemetry for application traces and synthetics.")
    elif t.startswith("For private OKE API endpoints"):
        replace_para_text(p,
            "Because the customer uses customer-managed Kubernetes, deployment must follow the documented manual Kubernetes Monitoring Solution pattern using Helm and/or Kubernetes manifests. Discovery must validate the supported Kubernetes version, cluster RBAC, namespaces, network/proxy/private egress, certificates, Management Agent/Gateway placement, resource requests/limits, upgrade and rollback procedure, and data residency. The platform team owns cluster access, upgrades, capacity and remediation; Observability & Security owns monitoring configuration, telemetry governance and alert quality.")

# Table 0: customer inputs. Remove ServiceNow row and correct Kubernetes consequence.
t0 = find_table(doc, "Input")
for row in list(t0.rows[1:]):
    if row.cells[0].text.strip() == "ServiceNow":
        remove_row(t0, row)
for row in t0.rows:
    if row.cells[0].text.strip() == "Containers":
        set_text(row.cells[2], "Manual Log Analytics Kubernetes Monitoring Solution deployment plus APM/OpenTelemetry and approved log/metric collection after compatibility, access and volume validation.")

# Table 1: in-scope service catalogue with service-description links beside each service.
t1 = find_table(doc, "OCI service")
for row in list(t1.rows[1:]):
    remove_row(t1, row)
catalogue = [
    ("OCI Monitoring", LINKS["monitoring"], "Metrics, alarms and service health", "Collects OCI and approved custom metrics; provides alarms and monitoring queries for infrastructure, routes and observability pipeline health.", "Baseline. Observability & Security defines alarm standards; product teams own thresholds and remediation."),
    ("OCI Logging", LINKS["logging"], "Central service/custom log ingestion and retention", "Receives approved OCI service and custom logs and supplies searchable retention before optional analytical routing.", "Baseline. Observability & Security owns log groups, retention and routing; product teams own log quality."),
    ("OCI Connector Hub", LINKS["connector"], "Managed telemetry routing", "Routes supported logs and service data to approved destinations such as Log Analytics and Notifications.", "Supporting component. No standalone SKU returned in the current pricing snapshot; downstream consumption remains chargeable."),
    ("OCI Log Analytics", LINKS["logan"], "Central log analytics and hybrid investigation", "Parses, enriches, clusters and investigates selected logs; provides dashboards, scheduled analytics and evidence links.", "Baseline for selected analytical logs; data volume, retention, residency, sources and parsers are discovery inputs."),
    ("Log Analytics Kubernetes Monitoring Solution", LINKS["k8s"], "Customer-managed Kubernetes visibility", "Collects selected Kubernetes logs, metrics and object state and provides cluster, workload, node and pod views.", "Manual Helm/manifests deployment. Validate supported version, RBAC, private egress, agents/gateways, resources, upgrades and rollback."),
    ("LoganAI", LINKS["loganai"], "AI-assisted log investigation", "Explains and summarizes logs, clusters and charts and suggests investigation follow-ups inside Log Analytics.", "Conditional after realm/region, IAM, commercial and residency review. Human validation required; it is not an agentic commander."),
    ("OCI Application Performance Monitoring", LINKS["apm"], "Application/API traces, topology and synthetics", "Collects APM/OpenTelemetry traces, maps service dependencies and runs approved synthetic journeys.", "Baseline for priority applications. APM active storage pricing is presumed through B95634 at one 300 GB active-storage unit, subject to quote confirmation."),
    ("OCI Management Agent and Management Gateway", LINKS["agent"], "Supported hybrid collection", "Provides the supported collection and communication path for selected cloud/on-premises targets and Log Analytics sources.", "Use where required. Validate proxies, ports, certificates, HA, upgrades and failure recovery. No external product agent is in scope."),
    ("OCI Notifications", LINKS["notifications"], "Approved alarm distribution", "Delivers approved Monitoring and Log Analytics alarms through supported HTTPS and email subscriptions.", "Baseline for alert delivery only. External incident/ticket workflow is outside this response."),
]
for name, url, fit, role, adoption in catalogue:
    row = t1.add_row()
    set_linked_cell(row.cells[0], name, url)
    set_text(row.cells[1], fit)
    set_text(row.cells[2], role)
    set_text(row.cells[3], adoption)
format_table(t1, [1.55, 1.7, 2.6, 2.15])

# Table 2: remove out-of-scope OEM/ServiceNow rows and keep answers in observability scope.
t2 = find_table(doc, "RFP section and requirement")
for row in list(t2.rows[1:]):
    req = " ".join(row.cells[0].text.split())
    if req.startswith("5.3 Use existing OEM") or req.startswith("6.4 Automatic ServiceNow"):
        remove_row(t2, row)
        continue
    ans = " ".join(row.cells[2].text.split())
    dep = " ".join(row.cells[3].text.split())
    replacements = [
        (r"\s*Keep existing tools and send material events to UA\.", " Existing product tools and downstream integrations remain outside this response."),
        (r"Proprietary OEM schemas still require UA adapters or partner integration\.", "Proprietary source formats require product-team interface support and may need a partner-built parser."),
        (r"Connector Hub and streaming routes", "Connector Hub routes"),
        (r"; Database Management covers supported database infrastructure", ""),
        (r" Business/customer impact still depends on UA topology and source data\.", " Business/customer impact depends on externally supplied service and customer context."),
        (r"; Group operational correlation remains in UA", "; cross-product correlation remains outside this response"),
        (r" UA performs cross-domain dedupe/correlation before ServiceNow policy\.", " Cross-product deduplication, correlation and ticketing are outside this response."),
        (r" and Operations Insights forecasts", ""),
        (r" UA supplies cross-domain correlation\.", " Cross-domain correlation is an external product-team responsibility."),
        (r" plus UA, OEM and ServiceNow context", " plus any separately approved external context"),
        (r"Object Storage", "Log Analytics archival storage"),
        (r"Streaming", "supported collection routes"),
        (r"Operations Insights can detect", "Log Analytics and APM can detect"),
        (r"Operations Insights forecasts supported database/host capacity; ", ""),
        (r"Other telecom forecasts require UA and/or governed OCI Data Science models\.", "Forecasting beyond available Monitoring metrics is outside this Observability response."),
        (r" include UA, collectors and product-team integrations", " include product-team platforms and external integrations"),
        (r"OKE components", "customer-managed Kubernetes monitoring components"),
        (r"OKE views", "customer-managed Kubernetes views"),
        (r", UA ticket enrichment", ""),
    ]
    for pat, repl in replacements:
        ans = re.sub(pat, repl, ans)
    dep = re.sub(r"Database-team onboarding and use-case accuracy acceptance\.", "Metric availability and use-case accuracy acceptance.", dep)
    set_text(row.cells[2], ans)
    set_text(row.cells[3], dep)
format_table(t2, [1.55, 0.7, 3.55, 2.1])

# Table 3: keep the LoganAI comparison, but remove named out-of-scope product interfaces.
t3 = find_table(doc, "Capability")
for row in t3.rows[1:]:
    cap = " ".join(row.cells[0].text.split())
    if cap.startswith("Reconcile UA topology"):
        set_text(row.cells[0], "Reconcile cross-domain topology and external workflow state")
        set_text(row.cells[1], "No")
        set_text(row.cells[2], "Custom requirement owned by the external AI/product workstream")
format_table(t3, [3.0, 1.4, 3.6])

# Table 4: OOTB gaps. Remove external product/OEM/database-management work and Pablo.
t4 = find_table(doc, "Gap")
for row in list(t4.rows[1:]):
    gap = " ".join(row.cells[0].text.split())
    if gap in {"Proprietary network/OEM collectors", "Database Management/Ops Insights onboarding"}:
        remove_row(t4, row)
        continue
    owner = " ".join(row.cells[2].text.split()).replace(", Operations and Security accept", " and Security accept").replace("Operations and Security acceptance", "Security acceptance")
    set_text(row.cells[2], owner)
format_table(t4, [1.65, 3.8, 2.45])

# Table 5: replace the platform BOM with only OCI Observability services and current public USD list prices.
t5 = find_table(doc, "Layer / component")
for row in list(t5.rows[1:]):
    remove_row(t5, row)
headers = ["OCI Observability service / description", "Part number and metric", "Current USD PAYG list price", "Sizing and scope note"]
for cell, value in zip(t5.rows[0].cells, headers):
    set_text(cell, value, bold=True)
bom = [
    ("OCI Monitoring — ingestion", LINKS["monitoring"], "B90925\nMillion Datapoints", "0–500: $0\n>500: $0.0025 per million datapoints", "Size from metric streams × dimensions × collection frequency. Tier bounds reproduce the pricing API."),
    ("OCI Monitoring — retrieval", LINKS["monitoring"], "B90926\nMillion Datapoints", "0–1,000: $0\n>1,000: $0.0015 per million datapoints", "Size dashboards, API queries, alarm evaluation and operator retrieval separately from ingestion."),
    ("OCI Logging — storage", LINKS["logging"], "B92593\nGigabyte Log Storage Per Month", "0–10 GB-month: $0\n>10: $0.05 per GB-month", "Applies to OCI Logging searchable storage. Avoid duplicate routing and unapproved payloads."),
    ("OCI Log Analytics — active storage", LINKS["logan"], "B95634\nLogging Analytics Storage Unit Per Month", "First 35 units: $372/unit-month\n>35–103: $260.40\n>103: $223.20", "API description states first 10 GB of log storage are free. Size selected analytical logs by active retention and filtering."),
    ("OCI Log Analytics — archival storage", LINKS["logan"], "B92809\nLogging Analytics Storage Unit Per Hour", "$0.02 per unit-hour", "Use only for the approved archive tier and retention policy; retrieval and rehydration assumptions require validation."),
    ("OCI Application Performance Monitoring", LINKS["apm"], "Commercial presumption: B95634\n300 GB active storage = 1 Log Analytics Active Storage Unit", "$372/month for the first assumed 300 GB unit (first-tier PAYG)", "Explicit proposal assumption, not a standalone APM SKU mapping returned by the current API. Confirm part number, unit conversion and rate in the final Oracle quote."),
    ("Log Analytics Kubernetes Monitoring Solution", LINKS["k8s"], "No standalone SKU in current API snapshot", "No separate list-price line; consumes Log Analytics active/archive storage and related telemetry services", "One manual Helm/manifests deployment per customer-managed Kubernetes cluster. Size logs/object state, nodes, pods, namespaces and retention."),
    ("OCI Management Agent and Management Gateway", LINKS["agent"], "No standalone SKU in current API snapshot", "No separate list-price line returned", "Required only where supported hybrid collection needs it; validate redundant gateways, proxy, certificates, resource footprint and lifecycle."),
    ("OCI Connector Hub", LINKS["connector"], "No standalone SKU in current API snapshot", "No separate list-price line returned; destination services remain chargeable", "Size connectors by source/destination/residency route and validate retries, duplicates, route health and replay."),
    ("OCI Notifications — HTTPS delivery", LINKS["notifications"], "B90940\nMillion Delivery Operations", "First 1 million: $0\n>1 million: $0.60 per million", "Used only for approved Observability alarm delivery endpoints; downstream incident processing is outside scope."),
    ("OCI Notifications — email delivery", LINKS["notifications"], "B90941\n1,000 Emails Sent", "First 1,000: $0\n>1,000: $0.02 per 1,000 emails", "Size approved operational subscriptions and suppression rules. SMS is not assumed."),
    ("LoganAI", LINKS["loganai"], "No standalone Observability SKU in current API snapshot", "No separate list-price line returned; validate any applicable GenAI consumption and service terms", "Analyst-assistance capability only. A custom GenAI coordinator and its AI/API SKUs belong to the OCI AI Team and are excluded from this BOM."),
]
for label, url, sku, price, note in bom:
    row = t5.add_row()
    set_linked_cell(row.cells[0], label, url)
    set_text(row.cells[1], sku)
    set_text(row.cells[2], price)
    set_text(row.cells[3], note)
format_table(t5, [1.9, 1.55, 1.75, 2.8])

# Table 6: OCI Observability implementation work only.
t6 = find_table(doc, "Work package")
for row in list(t6.rows[1:]):
    remove_row(t6, row)
work = [
    ("Discovery and sizing", "Inventory approved logs, custom metrics, Kubernetes clusters, priority applications/APIs, agents/gateways, regions and retention/residency classes; validate the APM 300 GB-unit commercial presumption.", "Measured ingestion, metric cardinality, spans, browser sessions, synthetics, clusters/nodes/pods, active/archive retention and Oracle quote confirmation."),
    ("OCI Observability foundation", "Configure Monitoring, Logging, Log Analytics, Connector Hub, Notifications and monitor-the-monitoring controls using the approved cloud/security foundation.", "Least privilege, data-class controls, route health, alarm ownership, non-production validation and rollback evidence."),
    ("Manual Kubernetes monitoring", "Deploy Log Analytics Kubernetes Monitoring Solution manually with Helm/manifests; configure supported collection, parsers, cluster identity and dashboards.", "Supported K8s version, RBAC, namespace scope, private egress/proxy, certificates, agent/gateway HA, resource limits, upgrade and rollback tests."),
    ("APM and OpenTelemetry", "Instrument priority applications/APIs and journeys; configure APM domains, service names, trace propagation, approved attributes, synthetics and SLO alarms.", "Trace continuity, sampling, sensitive-data controls, topology correctness, synthetic success/failure and 300 GB active-storage sizing assumption."),
    ("Log Analytics content", "Create and lifecycle-manage supported sources, parsers, enrichment, labels, dashboards, scheduled analytics, LoganAI workflows and evidence queries.", "Parsing accuracy, completeness/latency, false-positive review, access isolation and human validation."),
    ("PoV, resilience and handover", "Run load, failure, replay, agent/gateway outage, route, alert-delivery, data-residency and restore tests; provide runbooks, training and support handover.", "Signed acceptance metrics. Product teams accept access, instrumentation and remediation obligations for their own platforms."),
]
for vals in work:
    add_row(t6, vals)
format_table(t6, [1.6, 3.7, 2.7])

# Rename the BOM sections and replace the pricing assumptions / official references.
for p in doc.paragraphs:
    t = " ".join(p.text.split())
    if t == "7. Solution bill of materials (BOM)":
        replace_para_text(p, "7. OCI Observability bill of materials (BOM)")
    elif t.startswith("This BOM is a solution and procurement baseline"):
        replace_para_text(p,
            "This BOM is intentionally limited to OCI Observability services owned by the Observability & Security team. It excludes assurance, ITSM, management packs, database-management, networking, security-platform, storage/streaming, API integration and OCI Generative AI SKUs. Prices are public USD PAY_AS_YOU_GO list prices retrieved on 3 August 2026 from the Oracle List Pricing API snapshot last updated 16 July 2026. API tier bounds use rangeMin as exclusive and rangeMax as inclusive. Public list prices and the Cost Estimator are planning aids; the final Oracle quote/rate card, region, currency and contract govern.")
    elif t == "Platform and subscription BOM":
        replace_para_text(p, "OCI Observability subscription BOM and pricing metrics")
    elif t == "Implementation and professional-services BOM":
        replace_para_text(p, "OCI Observability implementation work packages")
    elif t == "Reference configuration assumptions":
        replace_para_text(p, "Commercial and sizing assumptions")

# Rewrite bullet assumptions after the BOM.
assumption_texts = [
    "APM commercial presumption: 1 × B95634 Log Analytics Active Storage Unit represents 300 GB of APM active storage. At the first public PAYG tier this is $372/month for the first assumed unit. Oracle must confirm the SKU/unit mapping in the final quote.",
    "Kubernetes is customer-managed, not OKE. Budget manual Helm/manifests deployment, cluster-specific compatibility and access validation, agent/gateway placement, upgrades and rollback for every approved cluster.",
    "No standalone SKU in the current API snapshot means only that no separate price line was returned; it does not promise that the component or its downstream consumption is free.",
    "Quantities remain TBD until measured ingestion, active/archive retention, metric cardinality, trace/session/synthetic volume, Kubernetes inventory, regions, non-production and DR requirements are approved.",
]
assumption_heading_idx = next(i for i,p in enumerate(doc.paragraphs) if " ".join(p.text.split()) == "Commercial and sizing assumptions")
for offset, text in enumerate(assumption_texts, 1):
    if assumption_heading_idx + offset < len(doc.paragraphs):
        p = doc.paragraphs[assumption_heading_idx + offset]
        if p.text.strip().startswith(("Production", "Retention", "All commercial", "APM")) or p.style.name.startswith("List"):
            replace_para_text(p, text)

# Official references: remove old cross-product list and retain implementation/pricing documentation.
ref_heading = next(p for p in doc.paragraphs if " ".join(p.text.split()) == "8. Official Oracle references")
replace_para_text(ref_heading, "8. Implementation documentation and supporting references")
ref_start = next(i for i, p in enumerate(doc.paragraphs) if p._p is ref_heading._p) + 1
for p in list(doc.paragraphs[ref_start:]):
    txt = " ".join(p.text.split())
    if txt.startswith("Connect with us"):
        break
    delete_paragraph(p)

refs = [
    ("Manual Kubernetes Monitoring Solution deployment and validation", LINKS["k8s"]),
    ("OCI APM implementation and instrumentation documentation", LINKS["apm"]),
    ("Management Agent and Management Gateway installation/operations", LINKS["agent"]),
    ("Log Analytics service implementation documentation", LINKS["logan"]),
    ("LoganAI usage and prerequisites", LINKS["loganai"]),
    ("OCI Monitoring alarms, metrics and queries", LINKS["monitoring"]),
    ("OCI Logging configuration and log management", LINKS["logging"]),
    ("OCI Connector Hub service-connector implementation", LINKS["connector"]),
    ("OCI Notifications topics and subscriptions", LINKS["notifications"]),
    ("Oracle List Pricing API access and field/tier semantics", LINKS["pricing_api_doc"]),
    ("Oracle public List Pricing API endpoint", LINKS["pricing_api"]),
    ("Oracle Cloud Cost Estimator", LINKS["cost_estimator"]),
]
anchor = ref_heading._p
for label, url in refs:
    p = doc.add_paragraph(style="List Bullet")
    p._p.getparent().remove(p._p)
    anchor.addnext(p._p)
    anchor = p._p
    add_hyperlink(p, label, url)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8.5)

# Final scrub: no Pablo, OKE, OEM or named assurance components remain.
for p in doc.paragraphs:
    if p._element is None:
        continue
    txt = p.text
    txt = re.sub(r"\bOKE\b", "customer-managed Kubernetes", txt)
    txt = txt.replace("Operations (Pablo)", "")
    txt = txt.replace("Pablo", "")
    txt = re.sub(r"\bOEM\b", "external product", txt)
    txt = re.sub(r"\bServiceNow\b", "external ITSM", txt)
    txt = re.sub(r"\bUA\b", "external assurance", txt)
    txt = txt.replace("Operations and Security acceptance", "Security acceptance")
    if txt != p.text:
        replace_para_text(p, txt)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = re.sub(r"\bOKE\b", "customer-managed Kubernetes", p.text)
                txt = txt.replace("Operations (Pablo)", "").replace("Pablo", "")
                txt = re.sub(r"\bOEM\b", "external product", txt)
                txt = re.sub(r"\bServiceNow\b", "external ITSM", txt)
                txt = re.sub(r"\bUA\b", "external assurance", txt)
                txt = txt.replace("Operations and Security acceptance", "Security acceptance")
                if txt != p.text:
                    replace_para_text(p, txt)

doc.core_properties.title = "AXIAN OCI Observability RFP Answers"
doc.core_properties.subject = "Scoped OCI Observability response and pricing BOM"
doc.core_properties.comments = "Version 1.4 — customer-managed Kubernetes and OCI Observability-only BOM"
doc.save(OUT)
print(OUT)
