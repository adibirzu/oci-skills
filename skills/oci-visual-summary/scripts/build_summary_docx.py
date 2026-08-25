#!/usr/bin/env python3
"""Create or append one accessible visual-summary page to a DOCX.

Uses the bundled python-docx workflow; generated files are scrubbed and audited
before returning.  The script never discovers a Python runtime or installs a
package: call it with the Python executable returned by the workspace loader.
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from urllib.parse import urlparse

from docx import Document  # python-docx, provided by the authoritative runtime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_STORYBOARD_SCENES_PER_PAGE = 4
_STORYBOARD_SERVICES_PER_PAGE = 8
_MAX_HANDOFF_BYTES = 8 * 1024 * 1024
_MAX_PREVIEW_BYTES = 8 * 1024 * 1024
_MAX_SOURCE_DOCX_BYTES = 64 * 1024 * 1024


def _bounded_regular_file(path: Path, *, label: str, maximum: int) -> Path:
    """Reject symlinks, special files, and oversized direct builder inputs."""
    if path.is_symlink() or not path.is_file() or path.stat().st_size <= 0 or path.stat().st_size > maximum:
        raise ValueError(f"{label} must be a bounded regular local file")
    return path


def _load_bounded_handoff(path: Path) -> dict:
    _bounded_regular_file(path, label="handoff", maximum=_MAX_HANDOFF_BYTES)
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("handoff must be a JSON object")
    return payload


def _validate_preview(path: Path | None) -> Path | None:
    if path is None:
        return None
    _bounded_regular_file(path, label="preview", maximum=_MAX_PREVIEW_BYTES)
    if path.read_bytes()[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError("preview must be a PNG file")
    return path


def _require_runtime() -> tuple[Path, Path]:
    node = os.environ.get("RUNTIME_NODE")
    skill_dir = os.environ.get("DOCUMENTS_SKILL_DIR")
    if not node or not skill_dir:
        raise RuntimeError("workspace runtime is unavailable; set RUNTIME_NODE and DOCUMENTS_SKILL_DIR from codex_app__load_workspace_dependencies")
    return Path(node), Path(skill_dir)


def _mark_operation(node: Path, skill_dir: Path) -> None:
    marker = skill_dir / "container_tools" / "mark_artifact_operation_started.mjs"
    subprocess.run([str(node), str(marker), "--operation-kind", "create", "--expected-output-count", "1", "--output-format", "docx"], check=True)


def _set_font(run, *, size: int, bold: bool = False, color: str = "18202B") -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_alt_text(inline, alt_text: str) -> None:
    """Give the optional inline preview meaningful OOXML alt text."""
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", "Visual summary preview")


def _add_external_hyperlink(paragraph, label: str, url: str) -> None:
    """Append a genuine DOCX hyperlink relationship, not visible URL text alone."""
    parsed = urlparse(url)
    if parsed.scheme != "https" or not parsed.netloc:
        raise RuntimeError("source links must use an absolute HTTPS URL")
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    run_pr.append(color); run_pr.append(underline)
    text = OxmlElement("w:t"); text.text = label
    run.append(run_pr); run.append(text); hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _storyboard_long_description(handoff: dict) -> str:
    supplied = handoff.get("accessibility", {}).get("long_description")
    if isinstance(supplied, str) and supplied.strip():
        return supplied.strip()
    final = next((page for page in handoff.get("pages", []) if isinstance(page, dict) and page.get("role") == "at-a-glance"), {})
    scene_titles = ", ".join(str(scene.get("title", "capability")) for scene in final.get("scenes", []) if isinstance(scene, dict))
    service_names = ", ".join(str(service.get("display_name", "OCI service")) for service in final.get("services", []) if isinstance(service, dict))
    return f"At-a-glance visual summary. Capability flow: {scene_titles or 'not supplied'}. OCI service map: {service_names or 'not supplied'}. Evidence: {handoff.get('evidence_footer', 'source ledger')}."


def _set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    tc_pr.append(shade)


def _mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _storyboard_physical_pages(pages: list[dict]) -> list[dict]:
    """Mirror the renderer's physical storyboard expansion for DOCX ordering."""
    expanded: list[dict] = []
    for page in pages:
        audience_role = str(page.get("role", ""))
        scenes = [item for item in page.get("scenes", []) if isinstance(item, dict)]
        services = [item for item in page.get("services", []) if isinstance(item, dict)]
        if audience_role == "capability-scenes":
            for scene in scenes:
                expanded.append(
                    {
                        **page,
                        "role": f"capability-scenes-{scene.get('unit_id', len(expanded) + 1)}",
                        "audience_role": audience_role,
                        "page_number": 1,
                        "page_count": 1,
                        "scenes": [scene],
                        "services": [],
                    }
                )
            continue
        scene_chunks = [scenes[index:index + _STORYBOARD_SCENES_PER_PAGE] for index in range(0, len(scenes), _STORYBOARD_SCENES_PER_PAGE)] or [[]]
        service_chunks = [services[index:index + _STORYBOARD_SERVICES_PER_PAGE] for index in range(0, len(services), _STORYBOARD_SERVICES_PER_PAGE)] or [[]]
        page_count = max(len(scene_chunks), len(service_chunks))
        for index in range(page_count):
            expanded.append(
                {
                    **page,
                    "role": audience_role if page_count == 1 else f"{audience_role}-{index + 1}",
                    "audience_role": audience_role,
                    "page_number": index + 1,
                    "page_count": page_count,
                    "scenes": scene_chunks[index] if index < len(scene_chunks) else [],
                    "services": service_chunks[index] if index < len(service_chunks) else [],
                }
            )
    return expanded


def _load_preview_manifest(path: Path | None) -> list[dict]:
    if path is None:
        return []
    _bounded_regular_file(path, label="preview manifest", maximum=2 * 1024 * 1024)
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, list):
        raise ValueError("preview manifest must be a JSON list")
    previews: list[dict] = []
    for index, item in enumerate(payload, start=1):
        if not isinstance(item, dict):
            raise ValueError("preview manifest entries must be objects")
        preview_path = Path(str(item.get("path", "")))
        try:
            _validate_preview(preview_path)
        except ValueError as exc:
            raise ValueError(f"preview manifest entry {index} is invalid: {exc}") from exc
        previews.append(
            {
                "role": str(item.get("role", item.get("audience_role", "storyboard"))),
                "audience_role": str(item.get("audience_role", item.get("role", "storyboard"))),
                "page_number": int(item.get("page_number", 1)),
                "page_count": int(item.get("page_count", 1)),
                "path": preview_path,
                "alt_text": str(item.get("alt_text", item.get("title", "Storyboard preview"))),
            }
        )
    return previews


def _append_summary(document: Document, handoff: dict, preview: Path | None, preview_manifest: list[dict] | None = None) -> None:
    section = document.sections[-1]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    if document.paragraphs:
        document.add_page_break()
    headline = handoff.get("headline_zone", {})
    title = headline.get("title") or handoff.get("title") or "Project at a glance"
    takeaway = headline.get("takeaway") or handoff.get("takeaway") or "A source-grounded visual summary."
    title_p = document.add_paragraph()
    title_p.style = document.styles["Heading 1"]
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title)
    _set_font(title_run, size=26, bold=True)
    takeaway_p = document.add_paragraph()
    takeaway_p.paragraph_format.space_after = Pt(10)
    takeaway_run = takeaway_p.add_run(takeaway)
    _set_font(takeaway_run, size=12, color="3C4655")
    if handoff.get("concept") == "illo-storyboard-sequence-v1":
        _append_storyboard_summary(document, handoff, preview, preview_manifest)
        return
    if handoff.get("visual_style", {}).get("variant") == "canvas-story-map":
        _append_canvas_story_map(document, handoff, preview)
        return
    if preview:
        image_p = document.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline = image_p.add_run().add_picture(str(preview), width=Inches(6.7))
        _set_alt_text(inline, handoff.get("accessibility", {}).get("alt_text", f"Visual summary of {title}"))
        caption = document.add_paragraph("Visual story map: the same key anchors are provided below as accessible text.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs: _set_font(run, size=9, color="53606E")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Anchor", "What it means", "Evidence")
    _mark_header_row(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        _set_cell_shading(cell, "E6B9AE")
        for run in cell.paragraphs[0].runs: _set_font(run, size=10, bold=True)
    for item in handoff.get("clusters", [])[:8]:
        cells = table.add_row().cells
        cells[0].text = str(item.get("title", ""))
        cells[1].text = str(item.get("detail", ""))
        cells[2].text = str(item.get("evidence_class", "code-backed"))
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs: _set_font(run, size=9, color="3C4655")
    source_heading = document.add_paragraph()
    source_heading.style = document.styles["Heading 2"]
    source_run = source_heading.add_run("Sources and evidence")
    _set_font(source_run, size=13, bold=True)
    source = document.add_paragraph(f"Evidence: {handoff.get('evidence_footer', 'source ledger')}")
    for run in source.runs: _set_font(run, size=9, color="53606E")
    source_register = handoff.get("source_register", [])
    if isinstance(source_register, list):
        for item in source_register:
            if not isinstance(item, dict) or not item.get("url"):
                continue
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{item.get('title', 'Official Oracle documentation')}: {item['url']}")
            _set_font(run, size=8, color="53606E")


def _storyboard_role_label(role: str) -> str:
    labels = {
        "project-promise": "Project promise",
        "workflow": "Workflow",
        "capability-scenes": "Capability scenes",
        "oci-service-map": "OCI service map",
        "at-a-glance": "At a glance",
    }
    return labels.get(role, role.replace("-", " ").title())


def _service_rendering_semantics(service: dict) -> str:
    """State when a public-stencil selection is rendered as a neutral glyph."""
    if service.get("rendered_as") != "neutral-service-glyph" or not service.get("fallback_reason"):
        return ""
    return f"Rendered as: neutral-service-glyph ({service['fallback_reason']})"


def _append_storyboard_summary(
    document: Document,
    handoff: dict,
    preview: Path | None,
    preview_manifest: list[dict] | None = None,
) -> None:
    preview_index = {
        str(entry["role"]): entry for entry in (preview_manifest or [])
    }

    pages = handoff.get("pages", [])
    if not isinstance(pages, list):
        pages = []
    for page in _storyboard_physical_pages(pages):
        audience_role = str(page.get("audience_role", page.get("role", "storyboard")))
        preview_entry = preview_index.get(str(page.get("role", audience_role)))
        if preview_entry:
            image_p = document.add_paragraph()
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # The image is the primary visual for each storyboard role, not a
            # thumbnail: keep it within the printable width while maximizing
            # legibility of its composed scene and service stencils.
            inline = image_p.add_run().add_picture(str(preview_entry["path"]), width=Inches(6.7))
            _set_alt_text(inline, f"{_storyboard_role_label(audience_role)}: {preview_entry['alt_text']}")
            caption = document.add_paragraph(f"{preview_entry['alt_text']} preview.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _set_font(run, size=9, color="53606E")
        elif preview and not preview_manifest:
            image_p = document.add_paragraph()
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inline = image_p.add_run().add_picture(str(preview), width=Inches(6.7))
            _set_alt_text(inline, handoff.get("accessibility", {}).get("alt_text", "Storyboard visual summary preview"))
            caption = document.add_paragraph("Storyboard visual summary preview. Ordered page text and service mappings follow.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _set_font(run, size=9, color="53606E")

        role = audience_role
        heading = document.add_paragraph()
        heading.style = document.styles["Heading 2"]
        heading_label = _storyboard_role_label(role)
        if int(page.get("page_count", 1)) > 1:
            heading_label = f"{heading_label} ({int(page.get('page_number', 1))}/{int(page.get('page_count', 1))})"
        heading_run = heading.add_run(heading_label)
        _set_font(heading_run, size=13, bold=True)
        if page.get("takeaway"):
            takeaway = document.add_paragraph(str(page["takeaway"]))
            for run in takeaway.runs:
                _set_font(run, size=10, color="3C4655")
        for index, scene in enumerate(page.get("scenes", []), start=1):
            scene_p = document.add_paragraph()
            scene_p.paragraph_format.space_after = Pt(5)
            title_run = scene_p.add_run(f"{index}. {scene.get('title', 'Capability')}. ")
            _set_font(title_run, size=10, bold=True)
            detail_run = scene_p.add_run(str(scene.get("detail", "")))
            _set_font(detail_run, size=10, color="3C4655")
            evidence_p = document.add_paragraph()
            evidence_p.paragraph_format.left_indent = Inches(0.18)
            evidence_run = evidence_p.add_run(f"Evidence: {str(scene.get('evidence_class', 'code-backed')).upper()}")
            _set_font(evidence_run, size=8, bold=True, color="53606E")
        for service in page.get("services", []):
            service_p = document.add_paragraph()
            service_p.paragraph_format.left_indent = Inches(0.18)
            service_text = f"{service.get('display_name', 'OCI service')} ({service.get('mapping_type', 'mappingType')})"
            semantics = _service_rendering_semantics(service)
            if semantics:
                service_text += f" — {semantics}"
            service_run = service_p.add_run(service_text)
            _set_font(service_run, size=9, bold=True, color="18202B")

    source_heading = document.add_paragraph()
    source_heading.style = document.styles["Heading 2"]
    source_run = source_heading.add_run("Sources and evidence")
    _set_font(source_run, size=13, bold=True)
    source = document.add_paragraph(f"Evidence: {handoff.get('evidence_footer', 'source ledger')}")
    for run in source.runs:
        _set_font(run, size=9, color="53606E")
    for item in handoff.get("source_register", []):
        if not isinstance(item, dict) or not item.get("url"):
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        prefix = paragraph.add_run(f"{item.get('title', 'Official Oracle documentation')}: ")
        _set_font(prefix, size=8, color="53606E")
        _add_external_hyperlink(paragraph, str(item["url"]), str(item["url"]))
    # Keep the accessible long description attached to the evidence block.
    # A separate trailing Heading 2 forced a nearly empty final page in short
    # storyboard companions even though the description itself is only a few
    # lines.  This compact native-text block remains searchable and editable.
    description = document.add_paragraph()
    description.paragraph_format.space_before = Pt(3)
    description.paragraph_format.space_after = Pt(0)
    description_label = description.add_run("Long description. ")
    _set_font(description_label, size=9, bold=True, color="18202B")
    description_run = description.add_run(_storyboard_long_description(handoff))
    _set_font(description_run, size=9, color="3C4655")


def _append_canvas_story_map(document: Document, handoff: dict, preview: Path | None) -> None:
    """Append a readable Canvas story-map companion without private art metadata.

    The canvas-story-map is a visual-first one-pager.  DOCX keeps one complete
    accessible preview (when supplied) and a native scene path so readers can
    understand and edit its content without relying on a flattened image.
    """
    canvas_layout = handoff.get("canvas_layout", {})
    # The public handoff carries only deterministic geometry; private prompts,
    # image receipts, and workstation paths never become DOCX text/metadata.
    _ = canvas_layout.get("composition", "scene-led")
    if preview:
        image_p = document.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline = image_p.add_run().add_picture(str(preview), width=Inches(6.7))
        _set_alt_text(inline, handoff.get("accessibility", {}).get("alt_text", "Canvas visual summary with an Oracle-red control thread and scene annotations"))
        caption = document.add_paragraph("Canvas visual summary. The editable scene path and evidence mapping follow.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            _set_font(run, size=9, color="53606E")
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 2"]
    run = heading.add_run("Scene path")
    _set_font(run, size=13, bold=True)
    for index, item in enumerate(handoff.get("clusters", [])[:8], start=1):
        scene = document.add_paragraph()
        scene.paragraph_format.space_after = Pt(5)
        title = str(item.get("title", "")).replace(f"{index}. ", "", 1)
        title_run = scene.add_run(f"Scene {index}: {title}. ")
        _set_font(title_run, size=10, bold=True)
        detail_run = scene.add_run(str(item.get("detail", "")))
        _set_font(detail_run, size=10, color="3C4655")
        services = item.get("service_names") or [item.get("service_label", "")]
        service_names = ", ".join(str(service) for service in services if service)
        if service_names:
            service_p = document.add_paragraph()
            service_p.paragraph_format.left_indent = Inches(0.18)
            service_run = service_p.add_run(f"OCI services: {service_names}")
            _set_font(service_run, size=9, color="C74634")
        evidence_p = document.add_paragraph()
        evidence_p.paragraph_format.left_indent = Inches(0.18)
        evidence_run = evidence_p.add_run(f"Evidence: {str(item.get('evidence_class', 'code-backed')).upper()}")
        _set_font(evidence_run, size=8, bold=True, color="53606E")
    source_heading = document.add_paragraph()
    source_heading.style = document.styles["Heading 2"]
    source_run = source_heading.add_run("Sources and evidence")
    _set_font(source_run, size=13, bold=True)
    source = document.add_paragraph(f"Evidence: {handoff.get('evidence_footer', 'source ledger')}")
    for run in source.runs:
        _set_font(run, size=9, color="53606E")
    for item in handoff.get("source_register", []):
        if not isinstance(item, dict) or not item.get("url"):
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{item.get('title', 'Official Oracle documentation')}: {item['url']}")
        _set_font(run, size=8, color="53606E")


def _run_tool(python: str, script: Path, args: list[str]) -> None:
    subprocess.run([python, str(script), *args], check=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--handoff", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--into", type=Path, help="copy an existing DOCX and append the one-page summary")
    parser.add_argument("--preview", type=Path, help="optional local PNG preview; key content remains text-equivalent")
    parser.add_argument("--preview-manifest", type=Path, help="optional JSON manifest of storyboard page previews")
    args = parser.parse_args()
    node, documents_skill = _require_runtime()
    handoff = _load_bounded_handoff(args.handoff)
    preview_manifest = _load_preview_manifest(args.preview_manifest)
    preview = _validate_preview(args.preview)
    if args.into:
        _bounded_regular_file(args.into, label="source DOCX", maximum=_MAX_SOURCE_DOCX_BYTES)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="visual-summary-docx-") as temp_dir:
        work = Path(temp_dir) / "authored.docx"
        if args.into:
            shutil.copy2(args.into, work)  # preserves the user source; it is never overwritten
            document = Document(work)
        else:
            document = Document()
        _mark_operation(node, documents_skill)  # exactly once, immediately before the edit/create
        _append_summary(document, handoff, preview, preview_manifest)
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.save(work)
        scrubbed = Path(temp_dir) / "scrubbed.docx"
        _run_tool(sys.executable, documents_skill / "scripts" / "privacy_scrub.py", [str(work), "--out", str(scrubbed)])
        _run_tool(sys.executable, documents_skill / "scripts" / "a11y_audit.py", [str(scrubbed), "--out_json", str(Path(temp_dir) / "a11y.json")])
        # render_docx.py is the mandatory all-page visual QA renderer.  Its
        # temporary PNGs are deliberately not copied beside the final DOCX.
        _run_tool(sys.executable, documents_skill / "render_docx.py", [str(scrubbed), "--output_dir", str(Path(temp_dir) / "rendered")])
        shutil.copy2(scrubbed, args.out)


if __name__ == "__main__":
    main()
